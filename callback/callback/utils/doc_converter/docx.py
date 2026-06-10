"""Markdown 转 DOCX"""

import re

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from callback.utils.doc_converter.common import parse_table


def _process_inline_formatting(paragraph, text: str):
    """处理行内格式（粗体、斜体、行内代码）"""
    parts = []
    last_end = 0

    inline_re = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")

    for match in inline_re.finditer(text):
        if match.start() > last_end:
            parts.append((text[last_end : match.start()], {}))

        if match.group(2):  # ***粗斜体***
            parts.append((match.group(2), {"bold": True, "italic": True}))
        elif match.group(3):  # **粗体**
            parts.append((match.group(3), {"bold": True}))
        elif match.group(4):  # *斜体*
            parts.append((match.group(4), {"italic": True}))
        elif match.group(5):  # `行内代码`
            parts.append((match.group(5), {"font_name": "Consolas", "font_size": 10}))

        last_end = match.end()

    if last_end < len(text):
        parts.append((text[last_end:], {}))

    if not parts:
        paragraph.add_run(text)
    else:
        for text_part, formatting in parts:
            if text_part:
                run = paragraph.add_run(text_part)
                if formatting.get("bold"):
                    run.bold = True
                if formatting.get("italic"):
                    run.italic = True
                if formatting.get("font_name"):
                    run.font.name = formatting["font_name"]
                if formatting.get("font_size"):
                    run.font.size = Pt(formatting["font_size"])


def _add_table_to_doc(doc, table_data):
    """添加表格到 Word 文档"""
    if not table_data:
        return

    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = "Table Grid"

    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                if i == 0:  # 表头加粗
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def _set_heading_black(heading):
    """强制设置标题文本颜色为黑色（移除默认主题色）"""
    for run in heading.runs:
        run.font.name = "Microsoft YaHei"
        run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._element.get_or_add_rPr()
        theme_color = rPr.find(qn("w:themeColor"))
        if theme_color is not None:
            rPr.remove(theme_color)


def markdown_to_docx(md_content: str) -> Document:
    """
    将 Markdown 转换为带格式的 Word 文档

    支持：标题、粗体/斜体/行内代码、无序/有序列表、表格、代码块、引用、分隔线
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    # 设置所有标题样式为黑色（移除默认主题色）
    for i in range(1, 10):
        try:
            heading_style = doc.styles[f"Heading {i}"]
            rPr = heading_style.element.get_or_add_rPr()
            theme_color = rPr.find(qn("w:themeColor"))
            if theme_color is not None:
                rPr.remove(theme_color)
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "000000")
            rPr.append(color)
        except KeyError:
            break

    lines = md_content.split("\n")
    i = 0
    prev_was_list = False

    while i < len(lines):
        line = lines[i]

        # 空行：跳过
        if not line.strip():
            prev_was_list = False
            i += 1
            continue

        # 标题处理
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 9)
            text = line.lstrip("#").strip()
            heading = doc.add_heading(text, level=level)
            _set_heading_black(heading)
            prev_was_list = False
            i += 1
            continue

        # 无序列表
        if line.strip().startswith(("- ", "* ", "+ ")):
            indent_level = (len(line) - len(line.lstrip())) // 2
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            if indent_level > 0:
                p.paragraph_format.left_indent = Pt(21 * (indent_level + 1))
            _process_inline_formatting(p, text)
            prev_was_list = True
            i += 1
            continue

        # 有序列表
        ordered_match = re.match(r"^(\d+)\.\s+(.*)$", line.strip())
        if ordered_match:
            list_num = int(ordered_match.group(1))
            list_text = ordered_match.group(2)
            indent_level = (len(line) - len(line.lstrip())) // 2

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(21 * (indent_level + 1))
            p.paragraph_format.first_line_indent = Pt(-21)

            num_run = p.add_run(f"{list_num}. ")
            num_run.font.name = "Microsoft YaHei"
            num_run.font.color.rgb = RGBColor(0, 0, 0)

            _process_inline_formatting(p, list_text)
            prev_was_list = True
            i += 1
            continue

        # 表格处理
        if "|" in line and i + 1 < len(lines) and re.match(
            r"^[\|\s\-:]+$", lines[i + 1].strip()
        ):
            table_data, rows_consumed = parse_table(lines, i)
            _add_table_to_doc(doc, table_data)
            i += rows_consumed
            prev_was_list = False
            continue

        # 代码块
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # 跳过结束的 ```

            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F5F5F5")
            shading.set(qn("w:val"), "clear")
            p.paragraph_format.element.get_or_add_pPr().append(shading)
            prev_was_list = False
            continue

        # 引用块
        if line.strip().startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_text = lines[i].strip()[1:].strip()
                if quote_text:
                    quote_lines.append(quote_text)
                i += 1
            quote_text = "\n".join(quote_lines)
            p = doc.add_paragraph(style="Quote")
            _process_inline_formatting(p, quote_text)
            prev_was_list = False
            continue

        # 分隔线
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.add_run("─" * 50)
            prev_was_list = False
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _process_inline_formatting(p, line)
        prev_was_list = False
        i += 1

    return doc
